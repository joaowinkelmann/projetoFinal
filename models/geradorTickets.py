from docx import Document
from docx.shared import Inches
import hashlib
import qrcode
from Crypto.Cipher import AES
from Crypto.Util.Padding import pad
from Crypto.Util.strxor import strxor
from models.gerenciadorProdutos import GerenciadorDeProdutos
from models.salvaArquivo import QRLog
from PIL import Image
from io import BytesIO

class SalesTicketGenerator:
    def __init__(self):
        self.gerenciador = GerenciadorDeProdutos()
        self.secret_key = b'mysecretkey12345'  # Replace with your own secret key

    def generate_sales_tickets(self, codigo, quantity):
        produto = self.gerenciador.buscar_produto(codigo)
        if not produto:
            print("Product not found.")
            return

        if quantity <= 0 or quantity > (produto.estoque - produto.vendido):
            print("Invalid quantity.")
            return

        document = Document()

        # Add header
        document.add_heading("Sales Tickets", 0)

        for index in range(1, quantity+1):
            # Add product details
            document.add_heading(f"Ticket #{index}", level=1)
            document.add_paragraph(f"Product Name: {produto.nome}")
            document.add_paragraph(f"Price: ${produto.preco_venda}")
            document.add_paragraph(f"Quantity: 1")
            document.add_paragraph(f"Total Price: ${produto.preco_venda}")

            # Generate unique identifier using truncated SHA-256
            identifier = self.generate_identifier(codigo, index)

            # Add QR code with the identifier
            document.add_heading("QR Code", level=2)
            qr_code_image = self.generate_qr_code(identifier)
            document.add_picture(qr_code_image, width=Inches(2), height=Inches(2))

        # Save the document
        filename = f"SalesTickets_{produto.nome.replace(' ', '_')}.docx"
        document.save(filename)
        print(f"Sales tickets saved to {filename}")

    def generate_identifier(self, product_id, index):
        identifier = f"{product_id}-{index}"
        sha256_hash = hashlib.sha256(identifier.encode()).hexdigest()
        truncated_hash = sha256_hash[:7]  # Truncate to 8 characters
        return truncated_hash

    def generate_qr_code(self, data):
        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_H, box_size=10, border=4)
        qr.add_data(data)
        qr.make(fit=True)
        qr_code_image = qr.make_image(fill_color="black", back_color="white")

        # Convert the image to bytes and create an in-memory image
        image_stream = BytesIO()
        qr_code_image.save(image_stream, format='PNG')
        image_stream.seek(0)

        return image_stream

